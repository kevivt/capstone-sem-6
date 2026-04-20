from __future__ import annotations

from pathlib import Path

import docx


def extract_docx_to_text(docx_path: Path) -> str:
    d = docx.Document(str(docx_path))
    lines: list[str] = []

    for para in d.paragraphs:
        t = (para.text or "").strip()
        if t:
            lines.append(t)

    for tbl in d.tables:
        for row in tbl.rows:
            row_text = [((cell.text or "").strip()) for cell in row.cells]
            row_text = [x for x in row_text if x]
            if row_text:
                lines.append(" | ".join(row_text))

    return "\n".join(lines).strip() + "\n"


def main() -> None:
    project_dir = Path(__file__).resolve().parent.parent
    templates_dir = project_dir / "esa templates" / "ESA - phase -2 report" / "ESA - phase -2 report" / "ESA - phase -2 report"
    out_dir = project_dir / "tmp" / "esa_template_extracted"
    out_dir.mkdir(parents=True, exist_ok=True)

    docx_files = sorted(templates_dir.glob("*.docx"))
    if not docx_files:
        raise FileNotFoundError(f"No .docx files found in: {templates_dir}")

    for p in docx_files:
        txt = extract_docx_to_text(p)
        out_path = out_dir / f"{p.stem}.txt"
        out_path.write_text(txt, encoding="utf-8")
        print(f"Extracted {p.name} -> {out_path}")


if __name__ == "__main__":
    main()

